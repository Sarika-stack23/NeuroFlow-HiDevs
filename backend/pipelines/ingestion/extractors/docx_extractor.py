"""DOCX extractor — paragraphs, tables, headings."""
from __future__ import annotations
from dataclasses import dataclass, field
from pipelines.ingestion.extractors.pdf_extractor import ExtractedPage


def extract_docx(file_path: str) -> list[ExtractedPage]:
    pages: list[ExtractedPage] = []
    try:
        from docx import Document
        doc = Document(file_path)

        current_section = ""
        current_heading = "Document"
        current_level = 0
        section_idx = 0

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            if style.startswith("Heading"):
                # Flush previous section
                if current_section.strip():
                    pages.append(ExtractedPage(
                        page_number=section_idx + 1,
                        content=current_section.strip(),
                        content_type="text",
                        metadata={"heading": current_heading, "heading_level": current_level},
                    ))
                    section_idx += 1
                try:
                    current_level = int(style.split()[-1])
                except ValueError:
                    current_level = 1
                current_heading = text
                current_section = ""
            else:
                current_section += text + "\n"

        # Flush last section
        if current_section.strip():
            pages.append(ExtractedPage(
                page_number=section_idx + 1,
                content=current_section.strip(),
                content_type="text",
                metadata={"heading": current_heading, "heading_level": current_level},
            ))

        # Tables
        for table in doc.tables:
            rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    rows.append("|" + "|".join(["---"] * len(cells)) + "|")
            pages.append(ExtractedPage(
                page_number=len(pages) + 1,
                content="\n".join(rows),
                content_type="table",
                metadata={},
            ))

    except Exception as e:
        pages.append(ExtractedPage(1, f"[DOCX error: {e}]", "text", {}))

    return pages
